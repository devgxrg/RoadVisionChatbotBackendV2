"""
Endpoints for the Analyze module.

Provides APIs for retrieving tender analysis results including:
- One-pager summaries
- Scope of work details
- RFP section analysis
- Data sheets
- Document templates
"""
from uuid import UUID
import logging
from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session

from app.db.database import get_db_session
from app.modules.analyze.db.schema import TenderAnalysis
from app.modules.analyze.models.pydantic_models import TenderAnalysisResponse
from app.modules.auth.services.auth_service import get_current_active_user
from app.modules.analyze.repositories import repository as analyze_repo
from app.modules.analyze.services import analysis_rfp_service as rfp_service
from app.modules.analyze.services import analysis_template_service as template_service

logger = logging.getLogger(__name__)
router = APIRouter()


@router.get(
    "/{tender_id}",
    response_model=TenderAnalysisResponse,
    summary="Get Tender Analysis",
    description="Retrieve analysis for a tender including one-pager, scope of work, RFP sections, datasheet, and templates. Returns available data regardless of completion status.",
    tags=["Analyze"],
)
def get_tender_analysis(
    tender_id: str,
    db: Session = Depends(get_db_session),
    current_user=Depends(get_current_active_user),
) -> TenderAnalysisResponse:
    """
    Retrieve the analysis for a tender, returning available data regardless of completion status.

    If analysis is still in progress, the response will include:
    - **status**: Current analysis status (pending, parsing, analyzing, completed, failed)
    - **progress**: Percentage complete (0-100)
    - **Completed fields**: Only populated once analysis reaches that stage
    - **Null fields**: Will be null if analysis hasn't reached that stage yet

    Returns:
    - **one_pager**: Executive summary with risk analysis and key highlights (available after analyzing stage)
    - **scope_of_work**: Detailed work packages, components, and deliverables (available after analyzing stage)
    - **rfp_sections**: Section-by-section RFP analysis with requirements and risks (optional, populate from AnalysisRFPSection table)
    - **data_sheet**: Key facts and figures in tabular format (available after analyzing stage)
    - **templates**: Document templates required for bidding (optional, populate from AnalysisDocumentTemplate table)

    Args:
        tender_id: The tender reference number (TDR) as string
        db: Database session
        current_user: Authenticated user

    Raises:
        HTTPException(404): If analysis not found

    Example:
        GET /api/v1/analyze/51655667

        Response while analyzing:
        {
            "id": "...",
            "tender_id": "51655667",
            "status": "analyzing",
            "progress": 45,
            "one_pager": null,
            "scope_of_work": null,
            "data_sheet": null,
            "rfp_sections": null,
            "templates": null
        }

        Response when completed:
        {
            "id": "...",
            "tender_id": "51655667",
            "status": "completed",
            "progress": 100,
            "analyzed_at": "2024-01-16T11:45:00Z",
            "one_pager": {...},
            "scope_of_work": {...},
            "data_sheet": {...},
            ...
        }
    """
    try:
        # Fetch the analysis record from database
        # Try both string and UUID to handle different input formats
        analysis = None
        try:
            analysis = analyze_repo.get_by_id(db, tender_id)
        except (ValueError, TypeError):
            # If tender_id is not a valid UUID, search by tender_id string
            analysis = db.query(TenderAnalysis).filter(
                TenderAnalysis.tender_id == tender_id
            ).first()

        if not analysis:
            logger.warning(f"Analysis not found for tender_id: {tender_id}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Analysis not found for tender {tender_id}",
            )

        rfp_section = rfp_service.get_rfp_sections(db, analysis.id)
        templates = template_service.get_templates(db, analysis.id)

        # Build the response with whatever data is available
        # Null fields indicate the analysis hasn't reached that stage yet
        response = TenderAnalysisResponse(
            id=str(analysis.id),
            tender_id=analysis.tender_id,
            status=analysis.status.value,
            progress=analysis.progress,
            analyzed_at=analysis.analysis_completed_at,
            # Include whatever analysis results are available
            # These will be None if analysis hasn't reached that stage
            one_pager=analysis.one_pager_json,
            scope_of_work=analysis.scope_of_work_json,
            data_sheet=analysis.data_sheet_json,
            # Fetch RFP sections and templates from related tables
            rfp_sections=rfp_section,
            templates=templates,
        )

        logger.info(
            f"Retrieved analysis for tender_id: {tender_id}, status: {analysis.status.value}, progress: {analysis.progress}"
        )
        return response

    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        logger.error(f"Error retrieving analysis for tender_id {tender_id}: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Error retrieving tender analysis",
        )


@router.post(
    "/trigger/{tender_ref}",
    summary="Trigger Tender Analysis",
    description="Start analysis for a tender by tender reference number. Creates analysis record and starts background processing.",
    tags=["Analyze"],
)
def trigger_tender_analysis(
    tender_ref: str,
    db: Session = Depends(get_db_session),
    # current_user=Depends(get_current_active_user),  # Removed auth for testing
):
    """
    Trigger analysis for a tender.
    
    Args:
        tender_ref: Tender reference number (e.g., "51184507")
        db: Database session
        current_user: Authenticated user
        
    Returns:
        Status message
    """
    try:
        from app.modules.analyze.scripts.analyze_tender import analyze_tender
        
        # Check if tender exists
        from app.modules.tenderiq.db.schema import Tender
        tender = db.query(Tender).filter(Tender.tender_ref_number == tender_ref).first()
        if not tender:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Tender {tender_ref} not found"
            )
        
        # Check if already analyzed
        existing = db.query(TenderAnalysis).filter(TenderAnalysis.tender_id == tender_ref).first()
        if existing and existing.status == AnalysisStatusEnum.COMPLETED:
            return {
                "status": "already_analyzed",
                "message": f"Tender {tender_ref} is already analyzed",
                "analysis_id": str(existing.id)
            }
        
        # Trigger analysis
        logger.info(f"Triggering analysis for tender {tender_ref}")
        analyze_tender(db, tender_ref)
        
        return {
            "status": "success",
            "message": f"Analysis triggered for tender {tender_ref}"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error triggering analysis for {tender_ref}: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error triggering analysis: {str(e)}"
        )


@router.get(
    "/report/download/{tender_id}",
    summary="Download Analysis Report",
    description="Download complete tender analysis report in PDF, Excel, or Word format",
    tags=["Analyze"],
)
def download_analysis_report(
    tender_id: str,
    format: str = "pdf",
    db: Session = Depends(get_db_session),
    current_user=Depends(get_current_active_user),
):
    """
    Download complete analysis report.

    Args:
        tender_id: Tender reference number
        format: Report format (pdf, excel, word) - defaults to pdf
        db: Database session
        current_user: Authenticated user

    Returns:
        File download response
    """
    from fastapi.responses import StreamingResponse
    import io

    try:
        # Fetch analysis data
        analysis = None
        try:
            analysis = analyze_repo.get_by_id(db, tender_id)
        except (ValueError, TypeError):
            analysis = db.query(TenderAnalysis).filter(
                TenderAnalysis.tender_id == tender_id
            ).first()

        if not analysis:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Analysis not found for tender {tender_id}"
            )

        # Fetch related data
        rfp_sections = rfp_service.get_rfp_sections(db, analysis.id)
        templates = template_service.get_templates(db, analysis.id)

        # Generate file based on format
        format_lower = format.lower()
        if format_lower in ['excel', 'xlsx']:
            file_content, filename, media_type = generate_excel_report(
                analysis, rfp_sections, templates
            )
        elif format_lower in ['word', 'docx']:
            file_content, filename, media_type = generate_word_report(
                analysis, rfp_sections, templates
            )
        else:
            # Default to PDF
            file_content, filename, media_type = generate_pdf_report(
                analysis, rfp_sections, templates
            )

        # Return file as streaming response
        return StreamingResponse(
            io.BytesIO(file_content),
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading report for {tender_id}: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error generating report: {str(e)}"
        )


@router.get(
    "/templates/download/{template_id}",
    summary="Download Template File",
    description="Generate and download a template file in the requested format (PDF, Excel, or Word)",
    tags=["Analyze"],
)
def download_template(
    template_id: str,
    db: Session = Depends(get_db_session),
    current_user=Depends(get_current_active_user),
):
    """
    Download a template file.
    
    Generates a blank template file based on the template metadata.
    Supports PDF, Excel (.xlsx), and Word (.docx) formats.
    
    Args:
        template_id: UUID of the template
        db: Database session
        current_user: Authenticated user
        
    Returns:
        File download response
    """
    from app.modules.analyze.db.schema import AnalysisDocumentTemplate
    from fastapi.responses import StreamingResponse
    import io
    
    try:
        # Fetch template from database
        template = db.query(AnalysisDocumentTemplate).filter(
            AnalysisDocumentTemplate.id == template_id
        ).first()
        
        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Template {template_id} not found"
            )
        
        # Determine format
        format_type = template.required_format.lower() if template.required_format else 'pdf'
        
        # Generate file based on format
        if format_type == 'pdf':
            file_content, filename, media_type = generate_pdf_template(template)
        elif format_type in ['excel', 'xlsx', 'xls']:
            file_content, filename, media_type = generate_excel_template(template)
        elif format_type in ['word', 'docx', 'doc']:
            file_content, filename, media_type = generate_word_template(template)
        else:
            # Default to PDF
            file_content, filename, media_type = generate_pdf_template(template)
        
        # Return file as streaming response
        return StreamingResponse(
            io.BytesIO(file_content),
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading template {template_id}: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error generating template: {str(e)}"
        )


def generate_pdf_template(template):
    """Generate a PDF template file"""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    import io
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1a56db'),
        spaceAfter=30,
    )
    story.append(Paragraph(template.template_name, title_style))
    story.append(Spacer(1, 0.2 * inch))
    
    # Description
    if template.description:
        desc_style = styles['BodyText']
        story.append(Paragraph(f"<b>Description:</b> {template.description}", desc_style))
        story.append(Spacer(1, 0.3 * inch))
    
    # Content Preview
    if template.content_preview:
        story.append(Paragraph("<b>Instructions:</b>", styles['Heading2']))
        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph(template.content_preview, styles['BodyText']))
        story.append(Spacer(1, 0.3 * inch))
    
    # Form fields placeholder
    story.append(Paragraph("<b>Form Fields:</b>", styles['Heading2']))
    story.append(Spacer(1, 0.1 * inch))
    
    # Create a simple table for form fields
    data = [
        ['Field Name', 'Value'],
        ['Company Name', ''],
        ['Address', ''],
        ['Contact Person', ''],
        ['Email', ''],
        ['Phone', ''],
        ['Date', ''],
    ]
    
    table = Table(data, colWidths=[3*inch, 4*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 10),
    ]))
    story.append(table)
    
    # Build PDF
    doc.build(story)
    
    filename = f"{template.template_name.replace(' ', '_')}.pdf"
    return buffer.getvalue(), filename, "application/pdf"


def generate_excel_template(template):
    """Generate an Excel template file"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    import io
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Template"
    
    # Title
    ws['A1'] = template.template_name
    ws['A1'].font = Font(size=16, bold=True, color="1a56db")
    ws.merge_cells('A1:D1')
    
    # Description
    if template.description:
        ws['A3'] = "Description:"
        ws['A3'].font = Font(bold=True)
        ws['B3'] = template.description
        ws.merge_cells('B3:D3')
    
    # Headers
    row_start = 6
    headers = ['Field Name', 'Value', 'Instructions', 'Mandatory']
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=row_start, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1a56db", end_color="1a56db", fill_type="solid")
        cell.alignment = Alignment(horizontal='center')
    
    # Sample fields
    fields = [
        ['Company Name', '', 'Enter your company name', 'Yes'],
        ['Address', '', 'Complete address', 'Yes'],
        ['Contact Person', '', 'Name of authorized person', 'Yes'],
        ['Email', '', 'Official email address', 'Yes'],
        ['Phone', '', 'Contact number', 'Yes'],
        ['Date', '', 'Submission date', 'Yes'],
    ]
    
    for idx, field_data in enumerate(fields, start=row_start+1):
        for col, value in enumerate(field_data, start=1):
            ws.cell(row=idx, column=col, value=value)
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 20
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 30
    ws.column_dimensions['D'].width = 15
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    filename = f"{template.template_name.replace(' ', '_')}.xlsx"
    return buffer.getvalue(), filename, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def generate_word_template(template):
    """Generate a Word template file"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    doc = Document()
    
    # Title
    title = doc.add_heading(template.template_name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(26, 86, 219)
    
    # Description
    if template.description:
        doc.add_paragraph()
        desc_para = doc.add_paragraph()
        desc_para.add_run('Description: ').bold = True
        desc_para.add_run(template.description)
    
    # Content Preview
    if template.content_preview:
        doc.add_paragraph()
        inst_heading = doc.add_heading('Instructions', level=2)
        doc.add_paragraph(template.content_preview)
    
    # Form Fields
    doc.add_paragraph()
    form_heading = doc.add_heading('Form Fields', level=2)
    
    # Create table
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Field Name'
    header_cells[1].text = 'Value'
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    fields = ['Company Name', 'Address', 'Contact Person', 'Email', 'Phone', 'Date']
    for idx, field in enumerate(fields, start=1):
        table.rows[idx].cells[0].text = field
        table.rows[idx].cells[1].text = ''
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"{template.template_name.replace(' ', '_')}.docx"
    return buffer.getvalue(), filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# ============================================================================
# ANALYSIS REPORT GENERATION FUNCTIONS
# ============================================================================

def generate_pdf_report(analysis, rfp_sections, templates):
    """Generate a comprehensive PDF report of the tender analysis"""
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    import io
    from datetime import datetime
    import html

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch
    )
    styles = getSampleStyleSheet()
    story = []

    # Custom styles with proper spacing
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a56db'),
        spaceAfter=30,
        spaceBefore=0,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold',
        leading=28
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#1a56db'),
        spaceAfter=16,
        spaceBefore=24,
        fontName='Helvetica-Bold',
        leading=20,
        keepWithNext=True
    )

    subheading_style = ParagraphStyle(
        'CustomSubHeading',
        parent=styles['Heading3'],
        fontSize=12,
        textColor=colors.HexColor('#333333'),
        spaceAfter=10,
        spaceBefore=14,
        fontName='Helvetica-Bold',
        leading=16,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=10,
        spaceAfter=8,
        spaceBefore=0,
        leading=14,
        alignment=TA_JUSTIFY,
        wordWrap='LTR'
    )
    
    # Cover Page
    story.append(Spacer(1, 2*inch))
    story.append(Paragraph("TENDER ANALYSIS REPORT", title_style))
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph(f"Tender ID: {analysis.tender_id}", styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    if analysis.analysis_completed_at:
        story.append(Paragraph(f"Analysis Date: {analysis.analysis_completed_at.strftime('%B %d, %Y')}", styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(f"Status: {analysis.status.value.upper()}", styles['Normal']))
    story.append(PageBreak())
    
    # ONE PAGER
    if analysis.one_pager_json:
        story.append(Paragraph("1. EXECUTIVE SUMMARY (ONE PAGER)", heading_style))
        one_pager = analysis.one_pager_json
        
        if one_pager.get('project_overview'):
            story.append(Paragraph("<b>Project Overview:</b>", subheading_style))
            story.append(Paragraph(one_pager['project_overview'], body_style))
            story.append(Spacer(1, 0.2*inch))

        if one_pager.get('financial_requirements'):
            story.append(Paragraph("<b>Financial Requirements:</b>", subheading_style))
            for req in one_pager['financial_requirements']:
                story.append(Paragraph(f"• {req}", body_style))
            story.append(Spacer(1, 0.2*inch))

        if one_pager.get('eligibility_highlights'):
            story.append(Paragraph("<b>Eligibility Highlights:</b>", subheading_style))
            for highlight in one_pager['eligibility_highlights']:
                story.append(Paragraph(f"• {highlight}", body_style))
            story.append(Spacer(1, 0.2*inch))

        if one_pager.get('important_dates'):
            story.append(Paragraph("<b>Important Dates:</b>", subheading_style))
            for date in one_pager['important_dates']:
                story.append(Paragraph(f"• {date}", body_style))
            story.append(Spacer(1, 0.2*inch))

        if one_pager.get('risk_analysis'):
            risk = one_pager['risk_analysis']
            story.append(Paragraph("<b>Risk Analysis:</b>", subheading_style))
            if risk.get('summary'):
                story.append(Paragraph(risk['summary'], body_style))
        
        story.append(PageBreak())
    
    # SCOPE OF WORK
    if analysis.scope_of_work_json:
        story.append(Paragraph("2. SCOPE OF WORK", heading_style))
        scope = analysis.scope_of_work_json
        
        if scope.get('project_details'):
            details = scope['project_details']
            story.append(Paragraph("<b>Project Details:</b>", subheading_style))
            details_data = []
            if details.get('project_name'):
                details_data.append(['Project Name', details['project_name']])
            if details.get('location'):
                details_data.append(['Location', details['location']])
            if details.get('duration'):
                details_data.append(['Duration', details['duration']])
            if details.get('contract_value'):
                details_data.append(['Contract Value', details['contract_value']])
            
            if details_data:
                table = Table(details_data, colWidths=[2*inch, 4.5*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ]))
                story.append(table)
                story.append(Spacer(1, 0.2*inch))
        
        if scope.get('work_packages'):
            story.append(Paragraph("<b>Work Packages:</b>", subheading_style))
            for i, package in enumerate(scope['work_packages'], 1):
                story.append(Paragraph(f"<b>{i}. {package.get('name', 'Work Package')}</b>", body_style))
                if package.get('description'):
                    story.append(Paragraph(package['description'], body_style))
                story.append(Spacer(1, 0.1*inch))
        
        story.append(PageBreak())
    
    # DATA SHEET
    if analysis.data_sheet_json:
        story.append(Paragraph("3. DATA SHEET", heading_style))
        datasheet = analysis.data_sheet_json
        
        sections = [
            ('Project Information', datasheet.get('project_information', [])),
            ('Contract Details', datasheet.get('contract_details', [])),
            ('Financial Details', datasheet.get('financial_details', [])),
            ('Technical Summary', datasheet.get('technical_summary', [])),
            ('Important Dates', datasheet.get('important_dates', []))
        ]
        
        for section_name, items in sections:
            if items:
                story.append(Paragraph(f"<b>{section_name}:</b>", subheading_style))
                data = [[item['label'], item['value']] for item in items]
                if data:
                    table = Table(data, colWidths=[2.5*inch, 4*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
                        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 0.2*inch))
        
        story.append(PageBreak())
    
    # RFP SECTIONS
    if rfp_sections and rfp_sections.sections:
        story.append(Paragraph("4. RFP SECTIONS ANALYSIS", heading_style))
        for section in rfp_sections.sections:
            story.append(Paragraph(f"<b>{section.section_name}: {section.section_title}</b>", subheading_style))
            if section.summary:
                story.append(Paragraph(section.summary, body_style))

            if section.key_requirements:
                story.append(Paragraph("<b>Key Requirements:</b>", body_style))
                for req in section.key_requirements[:5]:  # Limit to 5 for space
                    story.append(Paragraph(f"• {req}", body_style))

            story.append(Spacer(1, 0.15*inch))

        story.append(PageBreak())

    # TEMPLATES
    if templates:
        story.append(Paragraph("5. REQUIRED TEMPLATES", heading_style))
        all_templates = []
        all_templates.extend(templates.bid_submission_forms or [])
        all_templates.extend(templates.financial_formats or [])
        all_templates.extend(templates.technical_documents or [])
        all_templates.extend(templates.compliance_formats or [])

        if all_templates:
            story.append(Paragraph(f"<b>Total Templates: {len(all_templates)}</b>", body_style))
            story.append(Spacer(1, 0.1*inch))

            template_data = [['Template Name', 'Format', 'Mandatory']]
            for template in all_templates:
                template_data.append([
                    template.name,
                    template.format.upper(),
                    'Yes' if template.mandatory else 'No'
                ])

            table = Table(template_data, colWidths=[3.5*inch, 1.5*inch, 1.5*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a56db')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(table)

    # Build PDF
    doc.build(story)
    
    filename = f"Tender_Analysis_{analysis.tender_id}_{datetime.now().strftime('%Y%m%d')}.pdf"
    return buffer.getvalue(), filename, "application/pdf"


def generate_excel_report(analysis, rfp_sections, templates):
    """Generate a comprehensive Excel report of the tender analysis"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from datetime import datetime
    import io
    
    wb = Workbook()
    
    # Remove default sheet
    if 'Sheet' in wb.sheetnames:
        wb.remove(wb['Sheet'])
    
    # Header styles
    header_fill = PatternFill(start_color="1a56db", end_color="1a56db", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    title_font = Font(bold=True, size=16, color="1a56db")
    subheader_font = Font(bold=True, size=11)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # SUMMARY SHEET
    ws_summary = wb.create_sheet("Summary")
    ws_summary['A1'] = "TENDER ANALYSIS REPORT"
    ws_summary['A1'].font = title_font
    ws_summary.merge_cells('A1:D1')
    
    row = 3
    ws_summary[f'A{row}'] = "Tender ID:"
    ws_summary[f'A{row}'].font = subheader_font
    ws_summary[f'B{row}'] = analysis.tender_id
    row += 1
    
    ws_summary[f'A{row}'] = "Status:"
    ws_summary[f'A{row}'].font = subheader_font
    ws_summary[f'B{row}'] = analysis.status.value
    row += 1
    
    if analysis.analysis_completed_at:
        ws_summary[f'A{row}'] = "Analysis Date:"
        ws_summary[f'A{row}'].font = subheader_font
        ws_summary[f'B{row}'] = analysis.analysis_completed_at.strftime('%Y-%m-%d %H:%M')
        row += 1
    
    ws_summary.column_dimensions['A'].width = 20
    ws_summary.column_dimensions['B'].width = 40
    
    # ONE PAGER SHEET
    if analysis.one_pager_json:
        ws_one_pager = wb.create_sheet("One Pager")
        one_pager = analysis.one_pager_json
        row = 1
        
        ws_one_pager[f'A{row}'] = "EXECUTIVE SUMMARY"
        ws_one_pager[f'A{row}'].font = title_font
        ws_one_pager.merge_cells(f'A{row}:C{row}')
        row += 2
        
        if one_pager.get('project_overview'):
            ws_one_pager[f'A{row}'] = "Project Overview"
            ws_one_pager[f'A{row}'].font = subheader_font
            row += 1
            ws_one_pager[f'A{row}'] = one_pager['project_overview']
            ws_one_pager[f'A{row}'].alignment = Alignment(wrap_text=True)
            ws_one_pager.merge_cells(f'A{row}:C{row}')
            row += 2
        
        if one_pager.get('financial_requirements'):
            ws_one_pager[f'A{row}'] = "Financial Requirements"
            ws_one_pager[f'A{row}'].font = subheader_font
            row += 1
            for req in one_pager['financial_requirements']:
                ws_one_pager[f'A{row}'] = f"• {req}"
                row += 1
            row += 1
        
        ws_one_pager.column_dimensions['A'].width = 80
    
    # DATA SHEET
    if analysis.data_sheet_json:
        ws_datasheet = wb.create_sheet("Data Sheet")
        datasheet = analysis.data_sheet_json
        row = 1
        
        ws_datasheet[f'A{row}'] = "DATA SHEET"
        ws_datasheet[f'A{row}'].font = title_font
        ws_datasheet.merge_cells(f'A{row}:B{row}')
        row += 2
        
        sections = [
            ('Project Information', datasheet.get('project_information', [])),
            ('Contract Details', datasheet.get('contract_details', [])),
            ('Financial Details', datasheet.get('financial_details', [])),
            ('Technical Summary', datasheet.get('technical_summary', [])),
            ('Important Dates', datasheet.get('important_dates', []))
        ]
        
        for section_name, items in sections:
            if items:
                ws_datasheet[f'A{row}'] = section_name
                ws_datasheet[f'A{row}'].font = subheader_font
                ws_datasheet[f'A{row}'].fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
                ws_datasheet.merge_cells(f'A{row}:B{row}')
                row += 1
                
                for item in items:
                    ws_datasheet[f'A{row}'] = item['label']
                    ws_datasheet[f'B{row}'] = item['value']
                    ws_datasheet[f'A{row}'].border = border
                    ws_datasheet[f'B{row}'].border = border
                    row += 1
                row += 1
        
        ws_datasheet.column_dimensions['A'].width = 30
        ws_datasheet.column_dimensions['B'].width = 50

    # TEMPLATES SHEET
    if templates:
        ws_templates = wb.create_sheet("Templates")
        row = 1

        ws_templates[f'A{row}'] = "REQUIRED TEMPLATES"
        ws_templates[f'A{row}'].font = title_font
        ws_templates.merge_cells(f'A{row}:D{row}')
        row += 2

        # Headers
        headers = ['Template Name', 'Category', 'Format', 'Mandatory']
        for col, header in enumerate(headers, start=1):
            cell = ws_templates.cell(row=row, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
        row += 1

        # Data
        categories = [
            ('Bid Submission Forms', templates.bid_submission_forms or []),
            ('Financial Formats', templates.financial_formats or []),
            ('Technical Documents', templates.technical_documents or []),
            ('Compliance Formats', templates.compliance_formats or [])
        ]

        for category_name, items in categories:
            for template in items:
                ws_templates[f'A{row}'] = template.name
                ws_templates[f'B{row}'] = category_name
                ws_templates[f'C{row}'] = template.format.upper()
                ws_templates[f'D{row}'] = 'Yes' if template.mandatory else 'No'
                row += 1

        ws_templates.column_dimensions['A'].width = 50
        ws_templates.column_dimensions['B'].width = 25
        ws_templates.column_dimensions['C'].width = 15
        ws_templates.column_dimensions['D'].width = 15

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    filename = f"Tender_Analysis_{analysis.tender_id}_{datetime.now().strftime('%Y%m%d')}.xlsx"
    return buffer.getvalue(), filename, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def generate_word_report(analysis, rfp_sections, templates):
    """Generate a comprehensive Word report of the tender analysis"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime
    import io
    
    doc = Document()
    
    # Title Page
    title = doc.add_heading('TENDER ANALYSIS REPORT', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(26, 86, 219)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    info_para = doc.add_paragraph()
    info_para.add_run(f"Tender ID: {analysis.tender_id}\n").bold = True
    info_para.add_run(f"Status: {analysis.status.value}\n")
    if analysis.analysis_completed_at:
        info_para.add_run(f"Analysis Date: {analysis.analysis_completed_at.strftime('%B %d, %Y')}")
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ONE PAGER
    if analysis.one_pager_json:
        doc.add_heading('1. EXECUTIVE SUMMARY (ONE PAGER)', level=1)
        one_pager = analysis.one_pager_json
        
        if one_pager.get('project_overview'):
            doc.add_heading('Project Overview', level=2)
            doc.add_paragraph(one_pager['project_overview'])
        
        if one_pager.get('financial_requirements'):
            doc.add_heading('Financial Requirements', level=2)
            for req in one_pager['financial_requirements']:
                doc.add_paragraph(req, style='List Bullet')
        
        if one_pager.get('eligibility_highlights'):
            doc.add_heading('Eligibility Highlights', level=2)
            for highlight in one_pager['eligibility_highlights']:
                doc.add_paragraph(highlight, style='List Bullet')
        
        if one_pager.get('important_dates'):
            doc.add_heading('Important Dates', level=2)
            for date in one_pager['important_dates']:
                doc.add_paragraph(date, style='List Bullet')
        
        if one_pager.get('risk_analysis'):
            risk = one_pager['risk_analysis']
            doc.add_heading('Risk Analysis', level=2)
            if risk.get('summary'):
                doc.add_paragraph(risk['summary'])
        
        doc.add_page_break()
    
    # SCOPE OF WORK
    if analysis.scope_of_work_json:
        doc.add_heading('2. SCOPE OF WORK', level=1)
        scope = analysis.scope_of_work_json
        
        if scope.get('project_details'):
            details = scope['project_details']
            doc.add_heading('Project Details', level=2)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            if details.get('project_name'):
                row = table.add_row()
                row.cells[0].text = 'Project Name'
                row.cells[1].text = details['project_name']
            if details.get('location'):
                row = table.add_row()
                row.cells[0].text = 'Location'
                row.cells[1].text = details['location']
            if details.get('duration'):
                row = table.add_row()
                row.cells[0].text = 'Duration'
                row.cells[1].text = details['duration']
            if details.get('contract_value'):
                row = table.add_row()
                row.cells[0].text = 'Contract Value'
                row.cells[1].text = details['contract_value']
        
        if scope.get('work_packages'):
            doc.add_heading('Work Packages', level=2)
            for i, package in enumerate(scope['work_packages'], 1):
                doc.add_heading(f"{i}. {package.get('name', 'Work Package')}", level=3)
                if package.get('description'):
                    doc.add_paragraph(package['description'])
        
        doc.add_page_break()
    
    # DATA SHEET
    if analysis.data_sheet_json:
        doc.add_heading('3. DATA SHEET', level=1)
        datasheet = analysis.data_sheet_json
        
        sections = [
            ('Project Information', datasheet.get('project_information', [])),
            ('Contract Details', datasheet.get('contract_details', [])),
            ('Financial Details', datasheet.get('financial_details', [])),
            ('Technical Summary', datasheet.get('technical_summary', [])),
            ('Important Dates', datasheet.get('important_dates', []))
        ]
        
        for section_name, items in sections:
            if items:
                doc.add_heading(section_name, level=2)
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Grid Accent 1'
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Field'
                hdr_cells[1].text = 'Value'
                
                for item in items:
                    row = table.add_row()
                    row.cells[0].text = item['label']
                    row.cells[1].text = item['value']
        
        doc.add_page_break()
    
    # RFP SECTIONS
    if rfp_sections and rfp_sections.sections:
        doc.add_heading('4. RFP SECTIONS ANALYSIS', level=1)
        for section in rfp_sections.sections:
            doc.add_heading(f"{section.section_name}: {section.section_title}", level=2)
            if section.summary:
                doc.add_paragraph(section.summary)
            
            if section.key_requirements:
                doc.add_heading('Key Requirements', level=3)
                for req in section.key_requirements[:10]:
                    doc.add_paragraph(req, style='List Bullet')
        
        doc.add_page_break()

    # TEMPLATES
    if templates:
        doc.add_heading('5. REQUIRED TEMPLATES', level=1)
        all_templates = []
        all_templates.extend(templates.bid_submission_forms or [])
        all_templates.extend(templates.financial_formats or [])
        all_templates.extend(templates.technical_documents or [])
        all_templates.extend(templates.compliance_formats or [])

        if all_templates:
            doc.add_paragraph(f"Total Templates: {len(all_templates)}")

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Template Name'
            hdr_cells[1].text = 'Format'
            hdr_cells[2].text = 'Mandatory'

            for template in all_templates:
                row = table.add_row()
                row.cells[0].text = template.name
                row.cells[1].text = template.format.upper()
                row.cells[2].text = 'Yes' if template.mandatory else 'No'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"Tender_Analysis_{analysis.tender_id}_{datetime.now().strftime('%Y%m%d')}.docx"
    return buffer.getvalue(), filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
